import streamlit as st
from pdf2image import convert_from_bytes
import pytesseract
from docx import Document
import io

# ConfiguraÃ§Ã£o da PÃ¡gina
st.set_page_config(page_title="Conversor OCR", layout="centered")

def processar_pdf_para_docx(pdf_file):
    doc = Document()
    doc.add_heading('Texto ExtraÃ­do', 0)

    try:
        # Converter PDF para imagens (300 dpi Ã© bom para OCR)
        st.info("Lendo PDF... Isso pode demorar um pouco.")
        images = convert_from_bytes(pdf_file.read(), dpi=300)
        
        total = len(images)
        barra = st.progress(0)

        for i, image in enumerate(images):
            status_text = st.empty()
            status_text.text(f"Processando pÃ¡gina {i+1}/{total}...")
            
            if i > 0: doc.add_page_break()
            doc.add_heading(f'PÃ¡gina {i+1}', level=1)

            # OCR (MantÃ©m o config, mas processamos o resultado diferente)
            texto = pytesseract.image_to_string(image, lang='por+eng')
            
            if texto.strip():
                # NOVA LÃ“GICA: Divide o texto por linhas e cria parÃ¡grafos separados
                for linha in texto.split('\n'):
                    # SÃ³ adiciona se a linha nÃ£o for vazia
                    if linha.strip():
                        doc.add_paragraph(linha)
            else:
                doc.add_paragraph("[Sem texto detectado]")
            
            barra.progress((i + 1) / total)
            status_text.empty() # Limpa o texto de status

        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        return docx_io

    except Exception as e:
        st.error(f"Erro: {e}")
        return None

# Interface
st.title("ðŸ“„ Extrator de Texto (OCR)")
st.markdown("FaÃ§a upload de um PDF para receber um DOCX editÃ¡vel.")

uploaded_file = st.file_uploader("Arraste seu PDF aqui", type=['pdf'])

if uploaded_file and st.button("Processar Arquivo"):
    resultado = processar_pdf_para_docx(uploaded_file)
    if resultado:
        st.success("Pronto!")

        st.download_button("ðŸ“¥ Baixar Word (.docx)", resultado, "texto_convertido.docx")
